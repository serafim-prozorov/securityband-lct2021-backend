import os
import docx
import tf_fio
import xgboost as xgb



text_to_replace = 'Noname'
black_list_sym = '1234567890!№"#$%&()*+,-/:;<=>?@[\\]^_`{|}~\t\n'

class Docs():

    """Класс для обработки файлов docx"""
    def __init__(self, filename):
        self.filename = filename
        self.pagelist = []
        self.directory = '/files/'
        self.output_fname = ''
        self.doc = docx.Document(self.filename)
        self.model = tf_fio.Predictor()
        self.positive = 0
        self.find_tokens = 0


    def table_process(self):
        """Метод предназначен для анализа таблиц в файле docx,
        вызывает классификатор, прогнозирует сущности, и если в таблице содержатся
        персональные данные, заменяет их фразой Noname"""

        print('Обрабатываем таблицы...')
        for p in self.doc.tables:
            for t in p.table.rows:
                for c in t.cells:
                    splits = c.text.split(' ')
                    self.find_tokens = self.find_tokens+ len(splits)
                    predict = self.model.bst_predict(splits)
                    print(splits)
                    print(predict)
                    for i, p in enumerate(predict):

                        if (p == 1) and (splits[i] != '') and (splits[i] not in black_list_sym):
                            splits[i] = text_to_replace
                            self.positive += 1
                    str = ' '.join(splits)

                    c.text = str



    def paragraph_process(self):
        """Метод предназначен для анализа абзацев в файле docx,
        вызывает классификатор, прогнозирует сущности, и если в таблице содержатся
        персональные данные, заменяет их фразой Noname"""

        print('Обрабатываем абзацы...')
        for p in self.doc.paragraphs:
            for run in p.runs:
                splits = run.text.split(' ')
                self.find_tokens = self.find_tokens + len(splits)
                predict = self.model.bst_predict(splits)
                print(splits)
                print(predict)
                for i, p in enumerate(predict):
                    if (p == 1) and (splits[i]!='') and (splits[i] not in black_list_sym):
                        splits[i] =text_to_replace
                        self.positive += 1
                str = ' '.join(splits)
                run.text = str

    def report(self):
        """
        Отчет о количестве найденных сущностей
        """
                    #Negative                       #Positive          #Total
        return (self.find_tokens-self.positive, self.positive, self.find_tokens)




    def recogn_doc(self):
        """Метод сохраняет документ"""
        self.paragraph_process()
        self.table_process()
        self.output_fname = os.curdir + self.directory + 'out_' + self.filename.split("/")[-1]
        self.doc.save(self.output_fname)
        return self.output_fname



